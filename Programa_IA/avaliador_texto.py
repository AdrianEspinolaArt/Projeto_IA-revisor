from docx import Document
import spacy
from spellchecker import SpellChecker
from transformers import GPT2LMHeadModel, GPT2Tokenizer

class AvaliadorTexto:
    def __init__(self, perfil):
        self.perfil = perfil
        self.nlp = spacy.load('pt_core_news_lg')
        self.spell = SpellChecker(language='pt')
        self.gpt_model = GPT2LMHeadModel.from_pretrained("gpt2")
        self.gpt_tokenizer = GPT2Tokenizer.from_pretrained("gpt2")

    def avaliar_formatacao(self, doc):
        pontuacao = 0.5  # Pontuação inicial

        if doc.styles['Normal'].font.name.lower() != 'arial':
            pontuacao -= 0.1

        if doc.styles['Normal'].font.size.pt != 12:
            pontuacao -= 0.1

        if doc.styles['Normal'].paragraph_format.line_spacing != 1.5:
            pontuacao -= 0.1

        if doc.styles['Normal'].paragraph_format.first_line_indent.cm != 1.25:
            pontuacao -= 0.1

        if doc.styles['Normal'].paragraph_format.alignment != 3:  # Verifica se está justificado
            pontuacao -= 0.1

        return max(0, pontuacao)

    def avaliar_linhas(self, doc):
        pontuacao = 0.5  # Pontuação inicial

        num_linhas = len([p for p in doc.paragraphs if p.text.strip()])

        if 7 <= num_linhas <= 30:
            pontuacao += 0.5
        else:
            pontuacao -= 0.5

        return max(0, pontuacao)

    def avaliar_citacoes(self, doc):
        pontuacao = 1.0  # Pontuação inicial

        citacoes_presentes = any(
            ('"' in p.text and '(' in p.text and ')' in p.text) or
            ('(' in p.text and ')' in p.text)
            for p in doc.paragraphs
        )

        if not citacoes_presentes:
            pontuacao -= 1.0

        return max(0, pontuacao)

    def avaliar_lingua_portuguesa(self, doc):
        pontuacao = 2.0  # Pontuação inicial

        texto_completo = ' '.join(paragrafo.text for paragrafo in doc.paragraphs)
        doc_spacy = self.nlp(texto_completo)

        girias = ["cara", "mano", "pra", "tô"]
        if any(giria in texto_completo for giria in girias):
            pontuacao -= 0.2

        erros_ortografia = len(self.spell.unknown(texto_completo.split()))

        if erros_ortografia > 10:
            pontuacao -= 0.2
        elif erros_ortografia > 0:
            pontuacao -= 0.1

        erros_gramatica = sum(1 for token in doc_spacy if token.pos_ == 'X')

        if erros_gramatica > 10:
            pontuacao -= 0.2
        elif erros_gramatica > 0:
            pontuacao -= 0.1

        return max(0, pontuacao)

    def avaliar_adequacao(self, doc):
        pontuacao = 6.0  # Pontuação inicial

        for paragrafo in doc.paragraphs:
            # Utilize SpaCy para análise sintática e semântica
            doc_spacy = self.nlp(paragrafo.text)

            # Utilize o GPT para análise de coerência e clareza
            gpt_input = "Avalie a adequação do seguinte texto: " + paragrafo.text
            gpt_output = self.gpt_avaliar_coerencia(gpt_input)

            # Analise o resultado do GPT e ajuste a pontuação conforme necessário
            pontuacao_ajuste = self.analisar_gpt_output(gpt_output)
            pontuacao -= pontuacao_ajuste

            # Adapte a lógica para avaliação de adequação ao tema, atendimento à proposta,
            # adequação ao gênero, argumentação e exemplificação
            if 'tema' in paragrafo.text.lower():
                pontuacao += 1.0

            if 'atendimento à proposta' in paragrafo.text.lower():
                pontuacao += 1.0

            # Adicione mais lógica de avaliação conforme necessário

        return max(0, pontuacao)  # A pontuação não pode ser negativa

    def analisar_gpt_output(self, gpt_output):
        # Analise o resultado do GPT e ajuste a pontuação conforme necessário
        # Exemplo simples: contagem de palavras e frases coerentes
        num_palavras = len(gpt_output.split())
        num_frases_coerentes = gpt_output.count('.') + gpt_output.count('!') + gpt_output.count('?')

        pontuacao = 0.0

        # Ajuste da pontuação com base na contagem de palavras e frases
        if num_palavras > 50:
            pontuacao += 0.5

        if num_frases_coerentes > 3:
            pontuacao += 0.5

        return pontuacao

    def gpt_avaliar_coerencia(self, input_text):
        # Lógica para usar o GPT para avaliar coerência e clareza
        input_ids = self.gpt_tokenizer.encode(input_text, return_tensors="pt")
        output = self.gpt_model.generate(input_ids, max_length=100, num_beams=5, no_repeat_ngram_size=2, top_k=50, top_p=0.95, temperature=0.7)
        decoded_output = self.gpt_tokenizer.decode(output[0], skip_special_tokens=True)
        return decoded_output

    def avaliar_texto(self, caminho_arquivo):
        # Função principal que chama todas as funções de avaliação
        doc = Document(caminho_arquivo)
        pontuacao_formatacao = self.avaliar_formatacao(doc)
        pontuacao_linhas = self.avaliar_linhas(doc)
        pontuacao_citacoes = self.avaliar_citacoes(doc)
        pontuacao_lingua = self.avaliar_lingua_portuguesa(doc)
        pontuacao_adequacao = self.avaliar_adequacao(doc)

        # Adapte a lógica para incluir outras pontuações parciais

        pontuacao_total = pontuacao_formatacao + pontuacao_linhas + pontuacao_citacoes + pontuacao_lingua + pontuacao_adequacao

        return pontuacao_total

# Exemplo de uso:
avaliador = AvaliadorTexto(perfil="resenha")
caminho_do_arquivo = "caminho/do/seu/arquivo.docx"
pontuacao_final = avaliador.avaliar_texto(caminho_do_arquivo)
print(f"Pontuação final do texto: {pontuacao_final}")
