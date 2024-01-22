from docx import Document
import spacy
from autocorrect import Speller
from transformers import GPT2LMHeadModel, GPT2Tokenizer
import os
import enchant
import hunspell

class AvaliadorTexto:
    def __init__(self, perfil):
        self.perfil = perfil
        self.nlp = spacy.load('pt_core_news_lg')
        self.spell = Speller(lang='pt')
        self.gpt_model = GPT2LMHeadModel.from_pretrained("gpt2")
        self.gpt_tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
        
        # Caminho para o arquivo de gírias
        caminho_girias = 'girias.txt'

        # Carrega as gírias do arquivo
        self.GIRIAS = self.carregar_girias(caminho_girias)

        # Constantes
        self.PONTUACAO_INICIAL = {
            'Formatacao': 0.5,
            'Linhas': 0.5,
            'Citacoes': 1.0,
            'Lingua_Portuguesa': 2.0,
            'Adequacao': 6.0
        }

        self.ESTILOS = {
            'arial': 0.1,
            12: 0.1,
            1.5: 0.1,
            1.25: 0.1,
            3: 0.1
        }

        # Contadores de erros ortográficos e gramaticais
        self.erros_ortograficos = 0
        self.erros_gramaticais = 0 

    def carregar_girias(self, caminho):
        if os.path.exists(caminho):
            with open(caminho, 'r', encoding='utf-8') as file:
                return [linha.strip() for linha in file.readlines()]
        else:
            print(f"Arquivo de gírias não encontrado em {caminho}. Usando lista vazia de gírias.")
            return []    

    def verificar_estilo(self, normal_style):
        return sum(self.ESTILOS[estilo] if getattr(normal_style, estilo) != valor else 0
                   for estilo, valor in self.ESTILOS.items())

    def contar_erros_ortografia(self, texto_completo):
        doc_spacy = self.nlp(texto_completo)

        # Coleta palavras originais e corrige usando o Speller
        palavras_corrigidas = [self.spell(token.text) for token in doc_spacy]

        # Conjunto para armazenar palavras já avaliadas
        palavras_avaliadas = set()

        # Lista para armazenar palavras consideradas incorretas
        palavras_incorretas = []

        # Compara as palavras originais com as corrigidas para contar os erros
        erros_ortografia = 0
        for token, palavra_corrigida in zip(doc_spacy, palavras_corrigidas):
            # Verifica se a palavra já foi avaliada
            if token.text not in palavras_avaliadas:
                # Conta como um erro apenas se a palavra original e corrigida forem diferentes
                if token.text != palavra_corrigida:
                    erros_ortografia += 1
                    palavras_incorretas.append(token.text)

                # Adiciona a palavra ao conjunto de palavras avaliadas
                palavras_avaliadas.add(token.text)

        # Salva palavras incorretas em um arquivo
        with open('palavras_incorretas.txt', 'w', encoding='utf-8') as arquivo_saida:
            for palavra_incorreta in palavras_incorretas:
                arquivo_saida.write(f'{palavra_incorreta}\n')      

        return erros_ortografia

    def eh_erro_gramatical_avancado(self, token, doc):
        # Lógica para verificar se o token representa um possível erro gramatical
        if token.pos_ == 'VERB':
            # Verificar se há sujeito e objeto associados ao verbo
            sujeito_presente = any(token_dep.dep_ == 'nsubj' for token_dep in token.children)
            objeto_presente = any(token_dep.dep_ == 'obj' for token_dep in token.children)
            
            # Se não há sujeito ou objeto, considerar como erro
            if not (sujeito_presente and objeto_presente):
                return True

        # Adicione mais lógica conforme necessário para outras classes gramaticais

        return False

    def contar_erros_gramatica(self, doc_spacy):
        erros_gramatica = 0

        # Iterar sobre os tokens no documento
        for token in doc_spacy:
            # Lógica avançada para verificar se o token representa um possível erro gramatical
            if self.eh_erro_gramatical_avancado(token, doc_spacy):
                erros_gramatica += 1

        return erros_gramatica

    def avaliar_formatacao(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Formatacao']

        normal_style = getattr(doc.styles, 'Normal', None)
        if normal_style:
            pontuacao -= self.verificar_estilo(normal_style)

        return max(0, pontuacao)

    def avaliar_linhas(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Linhas']

        num_linhas = len([p for p in doc.paragraphs if p.text.strip()])

        if not (7 <= num_linhas <= 30):
            pontuacao -= 0.5

        return max(0, pontuacao)

    def avaliar_citacoes(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Citacoes']

        citacoes_presentes = any(
            ('"' in p.text and '(' in p.text and ')' in p.text) or
            ('(' in p.text and ')' in p.text)
            for p in doc.paragraphs
        )

        if not citacoes_presentes:
            pontuacao -= 1.0

        return max(0, pontuacao)

    def avaliar_lingua_portuguesa(self, doc):
        pontuacao_maxima = self.PONTUACAO_INICIAL['Lingua_Portuguesa']
        pontuacao = pontuacao_maxima

        texto_completo = ' '.join(paragrafo.text for paragrafo in doc.paragraphs)
        doc_spacy = self.nlp(texto_completo)

        # Avaliação de Gírias
        girias_presentes = any(giria in texto_completo for giria in self.GIRIAS)
        if girias_presentes:
            desconto_girias = 0.1
            pontuacao -= desconto_girias
            print(f"Desconto de {desconto_girias} por gírias aplicado. Pontuação após desconto: {pontuacao}")

        # Avaliação de Erros Ortográficos e Gramaticais
        erros_ortografia = self.contar_erros_ortografia(texto_completo)
        erros_gramatica = self.contar_erros_gramatica(doc_spacy)

        # Imprimir a quantidade de erros ortográficos antes de aplicar o desconto
        print(f"Número total de erros ortográficos encontrados: {erros_ortografia}")
        print(f"Número total de erros gramaticais encontrados: {erros_gramatica}")

        # Aqui, corrigimos os nomes das variáveis
        desconto_ortografia = self.aplicar_desconto_por_erros(erros_ortografia)
        desconto_gramatica = self.aplicar_desconto_por_erros(erros_gramatica)

        pontuacao -= desconto_ortografia + desconto_gramatica

        # Impressão de valores intermediários
        self.imprimir_valores_intermediarios(pontuacao_maxima, erros_ortografia, erros_gramatica)

        # Limita a pontuação ao valor máximo definido
        pontuacao = max(0, min(pontuacao, pontuacao_maxima))
        print("Pontuação final:", pontuacao)

        return pontuacao

    def aplicar_desconto_por_erros(self, num_erros):
        if 1 <= num_erros <= 10:
            return 0.1
        elif 11 <= num_erros <= 20:
            return 0.2
        elif num_erros == 0:
            return 0
        return 0

    def imprimir_valores_intermediarios(self, pontuacao_maxima, erros_ortografia, erros_gramatica):
        # Corrigido o nome da variável para pontuacao_maxima
        print("Pontuação antes dos descontos:", pontuacao_maxima)
        print("Desconto por erros ortográficos:", self.aplicar_desconto_por_erros(erros_ortografia))
        print("Desconto por erros gramaticais:", self.aplicar_desconto_por_erros(erros_gramatica))

    def avaliar_adequacao(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Adequacao']

        for paragrafo in doc.paragraphs:
            doc_spacy = self.nlp(paragrafo.text)

            gpt_input = "Avalie a adequacao do seguinte texto: " + paragrafo.text
            gpt_output = self.gpt_avaliar_coerencia(gpt_input)

            pontuacao_ajuste = self.analisar_gpt_output(gpt_output)
            pontuacao -= pontuacao_ajuste

            if 'tema' in paragrafo.text.lower():
                pontuacao += 1.0

            if 'atendimento a proposta' in paragrafo.text.lower():
                pontuacao += 1.0

        return max(0, pontuacao)

    def analisar_gpt_output(self, gpt_output):
        num_palavras = len(gpt_output.split())
        num_frases_coerentes = gpt_output.count('.') + gpt_output.count('!') + gpt_output.count('?')

        pontuacao = 0.0

        if num_palavras > 50:
            pontuacao += 0.5

        if num_frases_coerentes > 3:
            pontuacao += 0.5

        return pontuacao

    def gpt_avaliar_coerencia(self, input_text):
        input_ids = self.gpt_tokenizer.encode(input_text, return_tensors="pt")
        eos_token_id = self.gpt_tokenizer.eos_token_id
        print(f"ID do token de fim de sequencia: {eos_token_id}")
        
        output = self.gpt_model.generate(
            input_ids,
            max_length=263,
            num_beams=5,
            no_repeat_ngram_size=2,
            top_k=50,
            top_p=0.95,
            temperature=0.7,
            pad_token_id=50256
        )
        decoded_output = self.gpt_tokenizer.decode(output[0], skip_special_tokens=True)
        return decoded_output

    def imprimir_analise(self, pontuacoes):
        print("Analise de Pontuacoes:")
        print("----------------------")
        print("{:<25} {:<10}".format('Criterio', 'Pontuacao'))
        print("----------------------")

        for crit, pont in pontuacoes.items():
            print("{:<25} {:<10}".format(crit, pont))

        print("----------------------")
        print("Pontuacao Total: {:<10}".format(sum(pontuacoes.values())))

    def avaliar_texto(self, caminho_arquivo):
        doc = Document(caminho_arquivo)
        pontuacoes_detalhadas = {}

        for criterio in self.PONTUACAO_INICIAL:
            # Condição para evitar 'ERROS' como critério
            if criterio.lower() != 'erros':
                avaliacao_funcao = getattr(self, f'avaliar_{criterio.lower().replace(" ", "_").replace("ções", "coes").replace("ção", "cao")}')
                pontuacao = avaliacao_funcao(doc)

                # Adicione uma justificativa adequada para cada critério
                justificativa = self.obter_justificativa(criterio, pontuacao)

                pontuacoes_detalhadas[criterio] = {
                    'pontuacao': pontuacao,
                    'justificativa': justificativa
                }

        return pontuacoes_detalhadas

    def obter_justificativa(self, criterio, pontuacao):
        if criterio == 'Formatacao':
            if pontuacao >= 0.5:
                return 'A formatação está boa.'
            else:
                return 'Problemas na formatação.'
        elif criterio == 'Linhas':
            if pontuacao >= 0.5:
                return 'Número adequado de linhas.'
            else:
                return 'Número de linhas inadequado.'
        elif criterio == 'Citacoes':
            if pontuacao > 0.5:
                return 'Uso adequado de citações.'
            else:
                return 'Problemas com as citações.'
        elif criterio == 'Lingua_Portuguesa':
            if pontuacao > 1.5:
                return 'Bom uso da língua portuguesa.'
            else:
                return 'Problemas no uso da língua portuguesa.'
        elif criterio == 'Adequacao':
            if pontuacao > 3.0:
                return 'Texto adequadamente focado no tema.'
            else:
                return 'Problemas de adequação ao tema.'

        # Se nenhum critério correspondente for encontrado, retorna uma mensagem padrão
        return 'Justificativa não disponível para este critério.'