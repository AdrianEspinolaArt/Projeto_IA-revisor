from docx import Document
import spacy
from autocorrect import Speller
from transformers import GPT2LMHeadModel, GPT2Tokenizer
import os

class AvaliadorTexto:
    def __init__(self, perfil):
        self.perfil = perfil
        self.nlp = spacy.load('pt_core_news_lg')
        self.spell = Speller(lang='pt')
        self.gpt_model = GPT2LMHeadModel.from_pretrained("gpt2")
        self.gpt_tokenizer = GPT2Tokenizer.from_pretrained("gpt2")

        # Caminho para o arquivo de gírias
        caminho_girias = 'girias.txt'

        # Carrega as gírias do arquivo
        self.GIRIAS = self.carregar_girias(caminho_girias)

        # Constantes
        self.PONTUACAO_INICIAL = {
            'Formatacao': 0.5,
            'Linhas': 0.5,
            'Citacoes': 1.0,
            'Lingua_Portuguesa': 2.0,
            'Adequacao': 6.0
        }

        self.ESTILOS = {
            'arial': 0.1,
            12: 0.1,
            1.5: 0.1,
            1.25: 0.1,
            3: 0.1
        }

        # Contadores de erros ortográficos e gramaticais
        self.erros_ortograficos = 0
        self.erros_gramaticais = 0 

    def carregar_girias(self, caminho):
        if os.path.exists(caminho):
            with open(caminho, 'r', encoding='utf-8') as file:
                return [linha.strip() for linha in file.readlines()]
        else:
            print(f"Arquivo de gírias não encontrado em {caminho}. Usando lista vazia de gírias.")
            return []    

    def verificar_estilo(self, normal_style):
        return sum(self.ESTILOS[estilo] if getattr(normal_style, estilo) != valor else 0
                   for estilo, valor in self.ESTILOS.items())

    def contar_erros_ortografia(self, texto_completo):
        doc_spacy = self.nlp(texto_completo)

        # Coleta palavras originais e corrige usando o Speller
        palavras_corrigidas = [self.spell(token.text) if token.is_alpha and not token.text.isupper() else token.text for token in doc_spacy]

        # Conjunto para armazenar palavras já avaliadas
        palavras_avaliadas = set()

        # Conjunto para armazenar palavras consideradas incorretas
        palavras_incorretas = set()

        # Compara as palavras originais com as corrigidas para contar os erros
        erros_ortografia = 0
        for token, palavra_corrigida in zip(doc_spacy, palavras_corrigidas):
            # Verifica se a palavra já foi avaliada
            if token.text not in palavras_avaliadas:
                # Conta como um erro apenas se a palavra original e corrigida forem diferentes
                if token.text != palavra_corrigida:
                    erros_ortografia += 1
                    palavras_incorretas.add(token.text)

                # Adiciona a palavra ao conjunto de palavras avaliadas
                palavras_avaliadas.add(token.text)

        # Salva palavras incorretas em um arquivo
        with open('palavras_incorretas.txt', 'w', encoding='utf-8') as arquivo_saida:
            arquivo_saida.writelines(f'{palavra_incorreta}\n' for palavra_incorreta in palavras_incorretas)

        return erros_ortografia
    
    def ler_palavras_incorretas_do_arquivo(self):
        caminho_arquivo = 'palavras_incorretas.txt'
        palavras_incorretas = []

        try:
            with open(caminho_arquivo, 'r', encoding='utf-8') as arquivo:
                texto_original = arquivo.read()

                for linha in arquivo:
                    # Adiciona cada palavra incorreta à lista
                    palavra_incorreta = linha.strip()

                    # Corrige usando o Speller
                    palavra_corrigida = self.spell(palavra_incorreta)

                    # Encontra a posição da palavra no texto original
                    posicao_inicio = texto_original.find(palavra_incorreta)
                    posicao_fim = posicao_inicio + len(palavra_incorreta)

                    palavras_incorretas.append({
                        'original': palavra_incorreta,
                        'corrigida': palavra_corrigida,
                        'posicao_inicio': posicao_inicio,
                        'posicao_fim': posicao_fim
                    })
        except FileNotFoundError:
            print(f"Arquivo não encontrado em {caminho_arquivo}.")

        return palavras_incorretas

    def eh_erro_gramatical_avancado(self, token):
            # Lógica para verificar se o token representa um possível erro gramatical
            if token.pos_ == 'VERB':
                # Verificar se há sujeito associado ao verbo
                sujeito_presente = any(token_dep.dep_ == 'nsubj' for token_dep in token.children)
                
                # Se não há sujeito, considerar como erro
                if not sujeito_presente:
                    return True

            # Adicione mais lógica conforme necessário para outras classes gramaticais
            # Aqui adicionamos uma verificação para artigos indefinidos antes de substantivos
            if token.pos_ == 'DET' and token.text.lower() == 'um':
                proximo_token = token.head
                if proximo_token.pos_ == 'NOUN':
                    return True

            return False

    def contar_erros_gramatica(self, doc_spacy):
        erros_gramatica = 0
        palavras_incorretas = []

        # Iterar sobre os tokens no documento
        for token in doc_spacy:
            # Lógica avançada para verificar se o token representa um possível erro gramatical
            if self.eh_erro_gramatical_avancado(token):
                erros_gramatica += 1
                palavras_incorretas.append(token.text)

        # Salva palavras incorretas em um arquivo
        with open('palavras_incorretas_gramatica.txt', 'w', encoding='utf-8') as arquivo_saida:
            arquivo_saida.writelines(f'{palavra_incorreta}\n' for palavra_incorreta in palavras_incorretas)

        return erros_gramatica

    def avaliar_formatacao(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Formatacao']

        normal_style = getattr(doc.styles, 'Normal', None)
        if normal_style:
            pontuacao -= self.verificar_estilo(normal_style)

        return max(0, pontuacao)

    def avaliar_linhas(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Linhas']

        num_linhas = len([p for p in doc.paragraphs if p.text.strip()])

        if not (7 <= num_linhas <= 30):
            pontuacao -= 0.5

        return max(0, pontuacao)

    def avaliar_citacoes(self, doc):
        pontuacao = self.PONTUACAO_INICIAL['Citacoes']

        citacoes_presentes = any(
            ('"' in p.text and '(' in p.text and ')' in p.text) or
            ('(' in p.text and ')' in p.text)
            for p in doc.paragraphs
        )

        if not citacoes_presentes:
            pontuacao -= 1.0

        return max(0, pontuacao)
    def salvar_log(self, mensagem):
        with open('log.txt', 'a', encoding='utf-8') as arquivo_log:
            arquivo_log.write(mensagem + '\n')

    def avaliar_lingua_portuguesa(self, doc):
        pontuacao_maxima = self.PONTUACAO_INICIAL['Lingua_Portuguesa']
        pontuacao = pontuacao_maxima

        texto_completo = ' '.join(paragrafo.text for paragrafo in doc.paragraphs)
        doc_spacy = self.nlp(texto_completo)

        # Avaliação de Gírias
        girias_presentes = any(giria in texto_completo for giria in self.GIRIAS)
        if girias_presentes:
            desconto_girias = 0.1
            pontuacao -= desconto_girias
            print(f"Desconto de {desconto_girias} por gírias aplicado. Pontuação após desconto: {pontuacao}")

        # Avaliação de Erros Ortográficos e Gramaticais
        erros_ortografia = self.contar_erros_ortografia(texto_completo)
        erros_gramatica = self.contar_erros_gramatica(doc_spacy)

        # Imprimir a quantidade de erros ortográficos antes de aplicar o desconto
        print(f"Número total de erros ortográficos encontrados: {erros_ortografia}")
        print(f"Número total de erros gramaticais encontrados: {erros_gramatica}")

        # Aqui, corrigimos os nomes das variáveis
        desconto_ortografia = self.aplicar_desconto_por_erros(erros_ortografia)
        desconto_gramatica = self.aplicar_desconto_por_erros(erros_gramatica)

        pontuacao -= desconto_ortografia + desconto_gramatica

        # Exibir erros gramaticais
        self.contar_erros_gramatica(doc_spacy)

        # Impressão de valores intermediários
        self.imprimir_valores_intermediarios(pontuacao_maxima, erros_ortografia, erros_gramatica)
       
        mensagem_erros_ortografia = f"Número total de erros ortográficos encontrados: {erros_ortografia}"
        self.salvar_log(mensagem_erros_ortografia)

        # Imprimir a quantidade de erros gramaticais antes de aplicar o desconto
        mensagem_erros_gramatica = f"Número total de erros gramaticais encontrados: {erros_gramatica}"
        self.salvar_log(mensagem_erros_gramatica)

        # Limita a pontuação ao valor máximo definido
        pontuacao = max(0, min(pontuacao, pontuacao_maxima))
        print("Pontuação final:", pontuacao)

        return pontuacao

    def aplicar_desconto_por_erros(self, num_erros):
        if 1 <= num_erros <= 10:
            return 0.1
        elif 11 <= num_erros <= 60:
            return 0.2
        elif num_erros == 0:
            return 0
        return 0

    def imprimir_valores_intermediarios(self, pontuacao_maxima, erros_ortografia, erros_gramatica):
        # Corrigido o nome da variável para pontuacao_maxima
        print("Pontuação antes dos descontos:", pontuacao_maxima)
        print("Desconto por erros ortográficos:", self.aplicar_desconto_por_erros(erros_ortografia))
        print("Desconto por erros gramaticais:", self.aplicar_desconto_por_erros(erros_gramatica))

    def avaliar_adequacao(self, doc):
        # Simulando uma pontuação fictícia para testes
        pontuacao = 5.0

        return pontuacao

    def imprimir_analise(self, pontuacoes):
        print("Analise de Pontuacoes:")
        print("----------------------")
        print("{:<25} {:<10}".format('Criterio', 'Pontuacao'))
        print("----------------------")

        for crit, pont in pontuacoes.items():
            print("{:<25} {:<10}".format(crit, pont))

        print("----------------------")
        print("Pontuacao Total: {:<10}".format(sum(pontuacoes.values())))

    def avaliar_texto(self, caminho_arquivo):
        doc = Document(caminho_arquivo)
        pontuacoes_detalhadas = {}

        # Lê as palavras incorretas do arquivo
        palavras_incorretas = self.ler_palavras_incorretas_do_arquivo()

        for criterio in self.PONTUACAO_INICIAL:
            # Condição para evitar 'ERROS' como critério
            if criterio.lower() != 'erros':
                avaliacao_funcao = getattr(self, f'avaliar_{criterio.lower().replace(" ", "_").replace("ções", "coes").replace("ção", "cao")}')
                pontuacao = avaliacao_funcao(doc)

                # Adiciona uma justificativa adequada para cada critério
                justificativa = self.obter_justificativa(criterio, pontuacao)

                # Adiciona informações sobre palavras incorretas para a pontuação detalhada
                if criterio.lower() == 'lingua_portuguesa':
                    pontuacoes_detalhadas[criterio] = {
                        'pontuacao': pontuacao,
                        'justificativa': justificativa,
                        'palavras_incorretas': palavras_incorretas
                    }
                else:
                    pontuacoes_detalhadas[criterio] = {
                        'pontuacao': pontuacao,
                        'justificativa': justificativa
                    }

        return pontuacoes_detalhadas

    def obter_justificativa(self, criterio, pontuacao):
        if criterio == 'Formatacao':
            if pontuacao >= 0.5:
                return 'A formatação está boa.'
            else:
                return 'Problemas na formatação.'
        elif criterio == 'Linhas':
            if pontuacao >= 0.5:
                return 'Número adequado de linhas.'
            else:
                return 'Número de linhas inadequado.'
        elif criterio == 'Citacoes':
            if pontuacao > 0.5:
                return 'Uso adequado de citações.'
            else:
                return 'Problemas com as citações.'
        elif criterio == 'Lingua_Portuguesa':
            if pontuacao > 1.5:
                return 'Bom uso da língua portuguesa.'
            else:
                return 'Problemas no uso da língua portuguesa.'
        elif criterio == 'Adequacao':
            if pontuacao > 3.0:
                return 'Texto adequadamente focado no tema.'
            else:
                return 'Problemas de adequação ao tema.'

        # Se nenhum critério correspondente for encontrado, retorna uma mensagem padrão
        return 'Justificativa não disponível para este critério.'